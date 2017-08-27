__author__ = '@thebongy (Rishit Bansal) <rishit.bansal0@gmail.com>'
__license__ = 'MIT'
__status__ = 'Development'

import subprocess,os,sys
from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_UNDERLINE

FONT = 'Courier New'
SIZE = Pt(11)
TOP = Inches(0.5)
BOTTOM = Inches(1.5)
LEFT = Inches(0.7)
RIGHT = Inches(0.7)

INPUT_SNIPPET = '''

__oldRawInput__ = raw_input
def raw_input(prompt):
    result = __oldRawInput__(prompt+'\\n')
    print '{{DEBUG[INPUT]}}',result
    return result

'''

PRINT_SNIPPET = '''

import sys
_stdout = sys.stdout # Keeping a copy of sys.stdout

class OutStream(object):
    def __init__(self, target):
        self.target = target
    
    def write(self,s):
        self.target.write(s)
        self.target.flush() # Ensuring Output is always flushed!!

sys.stdout = OutStream(sys.stdout)

'''

if os.name == 'nt':
    PYTHON='C:\Python27\python.exe'
else:
    PYTHON='/usr/bin/python'
print 'Is your python executable located at',PYTHON,'?'

PYTHON = raw_input('If yes, just press enter, otherwise enter the exectuable path').strip() or PYTHON 
DIR = raw_input('Enter the directory with your assignment programs:\n')

FILES = raw_input('Enter filenames (without .py) in order, seperated by commas').split(',')

heading = raw_input('Enter the heading for the Outputs Word File:\n')

document = Document()
head_p = document.add_paragraph()
head_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
head_r = head_p.add_run(heading)

head_font = head_r.font
head_font.name = FONT
head_font.size = SIZE
head_font.bold = True
head_font.underline = WD_UNDERLINE.SINGLE

i = int(raw_input('Enter Starting Q no to show in output file'))


print '_______________________'
print '''Now the program will attempt to execute your files one by one. Note that
the output shown will be written to the word file simultaneously. If you make a mistake
while running a file, don't worry, you will be asked for the option of re-running the 
file in the end, and erasing your old output.

Notes:
1. Ignore the lines with {{DEBUG[INPUT]}}. These are used by the program
to format your input statements. These will not appear in the output file.
2. The current verion does not support the usage of the __future__module.
'''
raw_input('Press Enter to Continue.......')

command = PYTHON + ' ' + 'temp.py'

for f in FILES:
    print '____________________ Running ' + f + '.py'
    f_path = os.path.join(DIR,f+'.py')
    done = 'y'
    while done.lower() == 'y':
        temp = open('temp.py','w')
        temp.write(PRINT_SNIPPET)
        temp.write(INPUT_SNIPPET)
        temp.write(open(f_path).read())
        temp.close()

        proc = subprocess.Popen(command,
        shell=True,
        stdin = sys.stdin,
        stdout = subprocess.PIPE)
        output = ''

        while proc.poll() == None:
            data = proc.stdout.readline()
            print data.rstrip()
            output += data
        
        output+=proc.communicate()[0]
        print '____________________ File Execution Complete'
        done = raw_input('Would you like to rerun the file? (y for yes, n for no)')
    
    output = output.splitlines()
    prog_p = document.add_paragraph()
    prog_p.add_run(str(i)+')'+'\n').font.bold = True
    i += 1

    for line in output:
        bold = False
        if line.startswith('{{DEBUG[INPUT]}}'):
            bold = True
            line = line.replace('{{DEBUG[INPUT]}}','(Input)')
        line_r = prog_p.add_run(line + '\n')
        
        line_font = line_r.font
        
        if bold:
            line_font.bold = True
        
        line_font.name = FONT
        line_font.size = SIZE

sections = document.sections
for section in sections:
    section.top_margin = TOP
    section.bottom_margin = BOTTOM
    section.left_margin = LEFT
    section.right_margin = RIGHT
print 'ALL FILES HAVE BEEN RUN SUCCESFULLY!!'
x = raw_input('Enter the name to save the outputs word file as')
document.save(x+'.docx')
